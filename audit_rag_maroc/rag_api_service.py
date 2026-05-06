#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# ============================================================
# rag_api_service.py — Service FastAPI RAG Juridique Marocain
# Expose le RAG local via une API REST pour l'intégration
# avec le frontend Next.js du système d'audit
# ============================================================

import os
import sys
import io
import json
import uuid
from pathlib import Path
from datetime import datetime
from typing import Optional, List

# Forcer UTF-8 sur Windows
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

# python-docx pour la génération de rapports Word
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx non installé. Installez : pip install python-docx")

# Configuration du projet RAG
sys.path.insert(0, str(Path(__file__).parent))
from config import (
    VECTOR_DB_PATH, COLLECTION_NAME, EMBEDDING_MODEL,
    RETRIEVAL_K, AUDIT_PROMPT_TEMPLATE, AVAILABLE_MODELS, DEFAULT_MODEL
)

app = FastAPI(
    title="RAG Audit Juridique Marocain",
    description="API REST pour l'analyse juridique marocaine via RAG local",
    version="1.0.0"
)

# CORS pour permettre les appels depuis le frontend Next.js
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Modèles Pydantic ──────────────────────────────────────────

class AnalyseRequest(BaseModel):
    audit_id: str
    audit_title: str
    audit_description: Optional[str] = None
    document_texts: List[dict] = []   # [{"filename": "...", "content": "..."}]
    document_urls: List[dict] = []    # [{"filename": "...", "download_url": "...", "doc_id": "..."}]
    model: str = DEFAULT_MODEL

class AnalyseStatus(BaseModel):
    status: str  # "pending" | "running" | "done" | "error"
    result: Optional[dict] = None
    error: Optional[str] = None

# ── Cache en mémoire des analyses en cours/terminées ─────────
analyses_cache: dict[str, AnalyseStatus] = {}
reports_dir = Path(__file__).parent / "reports"
reports_dir.mkdir(exist_ok=True)

# ── Téléchargement et extraction de texte depuis URL ─────────

def download_and_extract_text(filename: str, url: str) -> str:
    """Télécharge un document depuis l'URL et extrait son texte."""
    try:
        import urllib.request
        import io

        print(f"📥 Téléchargement : {filename} depuis {url[:80]}...")
        req = urllib.request.Request(url, headers={'User-Agent': 'RAG-Audit/1.0'})
        with urllib.request.urlopen(req, timeout=30) as response:
            data = response.read()

        fname_lower = filename.lower()

        # PDF → extraction texte
        if fname_lower.endswith('.pdf'):
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(data))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                print(f"✅ PDF extrait : {len(text)} caractères ({len(reader.pages)} pages)")
                return text if text.strip() else f"[PDF vide ou non lisible : {filename}]"
            except Exception as e:
                print(f"⚠️ Erreur extraction PDF {filename}: {e}")
                return f"[Erreur extraction PDF : {e}]"

        # Word .docx → extraction texte
        elif fname_lower.endswith('.docx'):
            try:
                from docx import Document as DocxDoc
                doc = DocxDoc(io.BytesIO(data))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                print(f"✅ DOCX extrait : {len(text)} caractères")
                return text
            except Exception as e:
                return f"[Erreur extraction DOCX : {e}]"

        # Texte brut
        elif fname_lower.endswith(('.txt', '.csv')):
            return data.decode('utf-8', errors='replace')

        else:
            return f"[Format non supporté pour extraction texte : {filename}]"

    except Exception as e:
        print(f"❌ Erreur téléchargement {filename}: {e}")
        return f"[Erreur téléchargement : {e}]"


# ── Chargement du vectorstore ─────────────────────────────────

_vectorstore = None

def get_vectorstore():
    global _vectorstore
    if _vectorstore is not None:
        return _vectorstore
    db_path = Path(VECTOR_DB_PATH)
    if not db_path.exists():
        return None
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
        from langchain_community.vectorstores import Chroma
        embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
        _vectorstore = Chroma(
            persist_directory=str(db_path),
            embedding_function=embeddings,
            collection_name=COLLECTION_NAME,
        )
        count = _vectorstore._collection.count()
        print(f"✅ Base vectorielle chargée : {count} chunks")
        return _vectorstore
    except Exception as e:
        print(f"❌ Erreur chargement vectorstore : {e}")
        return None


# ── Analyse RAG ───────────────────────────────────────────────

def check_ollama_available(model_key: str) -> bool:
    """Vérifie si Ollama est disponible et le modèle est chargé."""
    try:
        import requests
        r = requests.get("http://localhost:11434/api/tags", timeout=3)
        if r.status_code == 200:
            models = [m["name"].split(":")[0] for m in r.json().get("models", [])]
            return model_key in models
    except Exception:
        pass
    return False


def run_rag_analysis(audit_id: str, request: AnalyseRequest):
    """Exécute l'analyse RAG en arrière-plan."""
    analyses_cache[audit_id] = AnalyseStatus(status="running")
    try:
        # ── Vérification 1 : des documents sont requis ────────────
        has_docs = bool(request.document_texts) or bool(request.document_urls)
        if not has_docs:
            result = {
                "audit_id": audit_id,
                "model_used": "N/A",
                "analysed_at": datetime.now().isoformat(),
                "risk_score": 0,
                "risk_level": "LOW",
                "violations": [],
                "sources": [],
                "no_documents": True,
                "summary": (
                    "⚠️ Aucun document n'a été uploadé sur cet audit. "
                    "L'analyse RAG nécessite au minimum un document (contrat, rapport, convention, etc.) "
                    "pour effectuer une comparaison avec la base juridique marocaine."
                ),
                "recommandations": [{
                    "action": "Uploadez les documents à analyser (contrats, rapports, conventions, statuts…) puis relancez l'analyse.",
                    "priorite": "IMMEDIATE",
                    "responsable": "Auditeur",
                    "reference": "",
                }],
                "conclusion": "Analyse impossible : veuillez d'abord uploader les documents de l'audit.",
            }
            if DOCX_AVAILABLE:
                report_path = generate_word_report(audit_id, request, result)
                result["report_path"] = str(report_path)
                result["report_filename"] = report_path.name
            analyses_cache[audit_id] = AnalyseStatus(status="done", result=result)
            return

        # ── Vérification 2 : Ollama disponible ? ──────────────────
        model_key = request.model if request.model in AVAILABLE_MODELS else DEFAULT_MODEL
        ollama_ok = check_ollama_available(model_key)

        if not ollama_ok:
            result = {
                "audit_id": audit_id,
                "model_used": "N/A",
                "analysed_at": datetime.now().isoformat(),
                "risk_score": 0,
                "risk_level": "LOW",
                "violations": [],
                "sources": [],
                "ollama_required": True,
                "summary": (
                    f"⚠️ Le modèle LLM '{model_key}' n'est pas disponible. "
                    "Pour effectuer une vraie analyse juridique, Ollama doit être installé et lancé sur votre PC "
                    f"avec le modèle '{model_key}' téléchargé."
                ),
                "recommandations": [
                    {
                        "action": "Installez Ollama depuis https://ollama.ai",
                        "priorite": "IMMEDIATE",
                        "responsable": "Administrateur",
                        "reference": "https://ollama.ai",
                    },
                    {
                        "action": f"Lancez Ollama puis téléchargez le modèle : ollama pull {model_key}",
                        "priorite": "IMMEDIATE",
                        "responsable": "Administrateur",
                        "reference": "",
                    },
                    {
                        "action": "Relancez l'analyse une fois Ollama démarré.",
                        "priorite": "COURT_TERME",
                        "responsable": "Auditeur",
                        "reference": "",
                    },
                ],
                "conclusion": f"Analyse bloquée : Ollama avec le modèle '{model_key}' est requis.",
            }
            if DOCX_AVAILABLE:
                report_path = generate_word_report(audit_id, request, result)
                result["report_path"] = str(report_path)
                result["report_filename"] = report_path.name
            analyses_cache[audit_id] = AnalyseStatus(status="done", result=result)
            return

        # ── Analyse RAG complète (Ollama + Documents disponibles) ──
        vectorstore = get_vectorstore()

        # Construction du texte à auditer depuis les documents
        texte_a_auditer = f"Mission d'audit : {request.audit_title}\n"
        if request.audit_description:
            texte_a_auditer += f"Description : {request.audit_description}\n"

        texte_a_auditer += "\n=== DOCUMENTS FOURNIS ===\n"

        # 1) Documents envoyés directement comme texte
        for doc in request.document_texts:
            fname = doc.get('filename', 'Document')
            content = doc.get('content', '').strip()
            if content and not content.startswith('[Contenu non disponible') and not content.startswith('[Erreur'):
                texte_a_auditer += f"\n--- {fname} ---\n"
                texte_a_auditer += content[:8000]
                if len(content) > 8000:
                    texte_a_auditer += "\n[... tronqué]"

        # 2) Documents à télécharger depuis les URLs
        for doc_url in request.document_urls:
            fname = doc_url.get('filename', 'Document')
            url   = doc_url.get('download_url', '')
            if url:
                content = download_and_extract_text(fname, url)
                if content and not content.startswith('[') :
                    texte_a_auditer += f"\n--- {fname} ---\n"
                    texte_a_auditer += content[:8000]
                    if len(content) > 8000:
                        texte_a_auditer += "\n[... tronqué]"
                else:
                    print(f"⚠️ Contenu non extrait pour {fname}: {content}")

        print(f"📝 Texte total à auditer : {len(texte_a_auditer)} caractères")

        violations = []
        sources_used = []

        if vectorstore:
            try:
                import requests as req_lib
                import json

                # ── Étape 1 : Récupérer les chunks juridiques pertinents ──
                print("🔍 Recherche des lois pertinentes dans la base vectorielle...")
                legal_chunks = vectorstore.similarity_search(texte_a_auditer, k=RETRIEVAL_K)

                # Collecter les sources
                seen = set()
                context_legal = ""
                for doc in legal_chunks:
                    nom = doc.metadata.get("nom_fichier", "Inconnu")
                    cat = doc.metadata.get("categorie", "")
                    page = doc.metadata.get("page", "?")
                    key = f"{nom}_p{page}"
                    context_legal += f"\n[SOURCE: {nom} | {cat} | page {page}]\n{doc.page_content}\n"
                    if key not in seen:
                        seen.add(key)
                        sources_used.append({
                            "fichier": nom,
                            "categorie": cat,
                            "page": str(page),
                            "extrait": doc.page_content[:200].replace("\n", " ")
                        })

                print(f"📚 {len(legal_chunks)} chunks juridiques récupérés")

                # ── Étape 2 : Construire le prompt et appeler Mistral directement ──
                prompt = f"""INSTRUCTIONS STRICTES : Tu es un expert juridique marocain. Réponds UNIQUEMENT en français. Utilise EXACTEMENT le format demandé ci-dessous. Ne fais PAS de commentaires généraux avant les violations.

TEXTES DE LOI MAROCAINS DISPONIBLES :
{context_legal}

DOCUMENT CLIENT À AUDITER :
{texte_a_auditer}

TÂCHE : Identifie TOUTES les violations du document client par rapport aux lois marocaines ci-dessus. Pour chaque problème trouvé dans le document, écris EXACTEMENT :

=== VIOLATION ===
📌 TEXTE ORIGINAL : [le texte exact du document client qui pose problème]
📖 CITATION EXACTE DU DOCUMENT : "[l'article ou la règle de loi marocaine qui est violée]"
📁 SOURCE : [le nom du fichier juridique source]

Si le document client est totalement conforme : écris uniquement ✅ CONFORME

REMARQUE IMPORTANTE : Commence directement par === VIOLATION === ou ✅ CONFORME. Ne fais PAS d'introduction.
"""

                print("🤖 Envoi à Mistral via Ollama...")
                resp = req_lib.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": model_key,
                        "prompt": prompt,
                        "stream": False,
                        "options": {"temperature": 0.0, "num_predict": 2048}
                    },
                    timeout=180
                )
                resp.raise_for_status()
                rag_answer = resp.json().get("response", "")
                print(f"✅ Mistral répondu : {len(rag_answer)} caractères")
                print(f"--- RÉPONSE MISTRAL ---\n{rag_answer[:500]}\n---")

                violations = parse_violations(rag_answer)
                model_used = f"RAG + Ollama ({model_key})"

            except Exception as e:
                raise RuntimeError(f"Erreur analyse: {e}")
        else:
            raise RuntimeError("Base vectorielle non chargée. Lancez d'abord 01_load_and_index.py")

        # Calcul du score de risque
        score = compute_risk_score(violations)
        level = score_to_level(score)

        result = {
            "audit_id": audit_id,
            "model_used": model_used,
            "analysed_at": datetime.now().isoformat(),
            "risk_score": score,
            "risk_level": level,
            "violations": violations,
            "sources": sources_used,
            "summary": build_summary(violations, score, level),
            "recommandations": build_recommendations(violations),
            "conclusion": build_conclusion(violations, score),
        }

        if DOCX_AVAILABLE:
            report_path = generate_word_report(audit_id, request, result)
            result["report_path"] = str(report_path)
            result["report_filename"] = report_path.name

        analyses_cache[audit_id] = AnalyseStatus(status="done", result=result)

    except Exception as e:
        analyses_cache[audit_id] = AnalyseStatus(status="error", error=str(e))


def parse_violations(rag_text: str) -> list:
    """Parse les violations depuis la réponse RAG de Mistral — parsing robuste."""
    print(f"🔎 Parsing réponse Mistral ({len(rag_text)} chars)...")

    # Cas conforme → aucune violation
    if "✅ CONFORME" in rag_text and "=== VIOLATION ===" not in rag_text:
        print("✅ Document conforme détecté")
        return []

    violations = []
    blocks = rag_text.split("=== VIOLATION ===")
    print(f"📋 Blocs de violations trouvés : {len(blocks) - 1}")

    for block in blocks[1:]:
        lines = block.strip().split("\n")
        violation = {
            "titre": "",
            "texte_original": "",
            "citation": "",
            "source": "",
            "severite": "MEDIUM",
        }
        current_field = None
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Détection des marqueurs (avec ou sans emoji)
            if "TEXTE ORIGINAL" in stripped or stripped.startswith("📌"):
                val = stripped.split(":", 1)[-1].strip().strip('"')
                violation["texte_original"] = val
                violation["titre"] = val[:120] + ("…" if len(val) > 120 else "")
                current_field = "texte"
            elif "CITATION" in stripped or stripped.startswith("📖"):
                val = stripped.split(":", 1)[-1].strip().strip('"').strip("'")
                violation["citation"] = val
                current_field = "citation"
            elif "SOURCE" in stripped or stripped.startswith("📁"):
                val = stripped.split(":", 1)[-1].strip().strip('"')
                violation["source"] = val
                current_field = "source"
            elif current_field == "texte" and not violation["texte_original"]:
                # Ligne de continuation si le champ est vide
                violation["texte_original"] = stripped
                violation["titre"] = stripped[:120]

        # Si le bloc contient du texte mais pas de marqueurs — tenter extraction brute
        if not violation["texte_original"] and not violation["citation"]:
            # Prendre le premier contenu non vide du bloc
            for line in lines:
                line = line.strip()
                if line and not line.startswith("==="):
                    violation["texte_original"] = line[:200]
                    violation["titre"] = line[:120]
                    break

        # Déduire la sévérité
        texte_combined = (violation["texte_original"] + " " + violation["citation"]).lower()
        if any(kw in texte_combined for kw in ["fraude", "pénal", "criminel", "nullité", "prison", "escroquerie"]):
            violation["severite"] = "CRITICAL"
        elif any(kw in texte_combined for kw in ["licenciement", "sanction", "amende", "infraction", "illégal", "interdit", "non déclaré", "non payé"]):
            violation["severite"] = "HIGH"
        elif any(kw in texte_combined for kw in ["délai", "procédure", "formalité", "obligation", "minimum", "salaire", "capital"]):
            violation["severite"] = "MEDIUM"
        else:
            violation["severite"] = "MEDIUM"  # MEDIUM par défaut (pas LOW)

        if violation["texte_original"] or violation["citation"]:
            violations.append(violation)
            print(f"  ✓ Violation: [{violation['severite']}] {violation['titre'][:60]}")

    print(f"📊 Total violations parsées : {len(violations)}")
    return violations


def compute_risk_score(violations: list) -> int:
    if not violations:
        return 0
    score = 0
    for v in violations:
        s = v.get("severite", "MEDIUM")
        score += {"CRITICAL": 30, "HIGH": 20, "MEDIUM": 10, "LOW": 5}.get(s, 10)
    return min(score, 100)


def score_to_level(score: int) -> str:
    if score >= 70: return "CRITICAL"
    if score >= 40: return "HIGH"
    if score >= 20: return "MEDIUM"
    return "LOW"


def build_summary(violations, score, level) -> str:
    n = len(violations)
    if n == 0:
        return "Aucune violation détectée dans les documents fournis. Le dossier semble conforme au droit marocain."
    return (f"L'analyse RAG juridique marocaine a détecté {n} point(s) nécessitant attention. "
            f"Score de risque : {score}/100 (niveau {level}). "
            f"Ces points sont comparés aux textes de loi marocains indexés dans la base documentaire.")


def build_recommendations(violations) -> list:
    recs = []
    for v in violations:
        sev = v.get("severite", "MEDIUM")
        recs.append({
            "action": f"Vérifier la conformité : {v.get('titre', '')}",
            "priorite": "IMMEDIATE" if sev == "CRITICAL" else "COURT_TERME" if sev == "HIGH" else "LONG_TERME",
            "responsable": "Auditeur",
            "reference": v.get("source", ""),
        })
    return recs


def build_conclusion(violations, score) -> str:
    if score < 20:
        return "Le dossier présente un faible niveau de risque juridique selon les textes marocains."
    if score < 50:
        return "Des points de vigilance ont été identifiés. Une révision ciblée est recommandée."
    return "Des non-conformités significatives ont été détectées. Une intervention urgente de l'auditeur est requise."


# ── Génération du rapport Word ────────────────────────────────

def generate_word_report(audit_id: str, request: AnalyseRequest, result: dict) -> Path:
    """Génère un rapport Word professionnel complet."""
    doc = DocxDocument()

    # Styles généraux
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Page de titre ──────────────────────────────────────────
    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_par.add_run("⚖️ RAPPORT D'ANALYSE JURIDIQUE\nRAG DROIT MAROCAIN")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # bleu

    doc.add_paragraph()

    # Sous-titre mission
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Mission : {request.audit_title}")
    sr.bold = True
    sr.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date d'analyse : {datetime.now().strftime('%d/%m/%Y à %H:%M')} | "
                 f"Modèle : {result.get('model_used', 'RAG')} | "
                 f"Réf. : AUDIT-{audit_id[:8].upper()}")

    doc.add_page_break()

    # ── Résumé exécutif ────────────────────────────────────────
    h1 = doc.add_heading("1. RÉSUMÉ EXÉCUTIF", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    score = result.get("risk_score", 0)
    level = result.get("risk_level", "LOW")
    colors = {"CRITICAL": "CC0000", "HIGH": "E65C00", "MEDIUM": "B8860B", "LOW": "006400"}
    level_labels = {"CRITICAL": "CRITIQUE", "HIGH": "ÉLEVÉ", "MEDIUM": "MODÉRÉ", "LOW": "FAIBLE"}

    score_p = doc.add_paragraph()
    score_p.add_run("Score de Risque Juridique : ").bold = True
    sr2 = score_p.add_run(f"{score}/100 — {level_labels.get(level, level)}")
    sr2.bold = True
    col = colors.get(level, "000000")
    sr2.font.color.rgb = RGBColor(int(col[0:2], 16), int(col[2:4], 16), int(col[4:6], 16))
    sr2.font.size = Pt(14)

    doc.add_paragraph(result.get("summary", ""))

    # ── Violations détectées ───────────────────────────────────
    doc.add_heading("2. NON-CONFORMITÉS DÉTECTÉES", level=1)

    violations = result.get("violations", [])
    if not violations:
        doc.add_paragraph("✅ Aucune non-conformité majeure détectée dans les documents fournis.")
    else:
        for i, v in enumerate(violations, 1):
            sev = v.get("severite", "MEDIUM")
            sev_labels = {"CRITICAL": "🔴 CRITIQUE", "HIGH": "🟠 ÉLEVÉ", "MEDIUM": "🟡 MODÉRÉ", "LOW": "🟢 FAIBLE"}

            viol_heading = doc.add_paragraph()
            r = viol_heading.add_run(f"Violation {i} : {v.get('titre', 'Sans titre')}")
            r.bold = True
            r.font.size = Pt(12)

            doc.add_paragraph(f"Sévérité : {sev_labels.get(sev, sev)}")

            if v.get("texte_original"):
                p = doc.add_paragraph()
                p.add_run("Texte concerné : ").bold = True
                p.add_run(v["texte_original"])

            if v.get("citation"):
                p2 = doc.add_paragraph()
                p2.add_run("Référence légale : ").bold = True
                p2.add_run(f'"{v["citation"]}"')
                p2.runs[-1].italic = True

            if v.get("source"):
                doc.add_paragraph(f"Source : {v['source']}")

            if v.get("categorie"):
                doc.add_paragraph(f"Catégorie : {v['categorie']}")

            doc.add_paragraph()

    # ── Sources juridiques utilisées ───────────────────────────
    sources = result.get("sources", [])
    if sources:
        doc.add_heading("3. SOURCES JURIDIQUES MAROCAINES CONSULTÉES", level=1)
        for s in sources:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{s.get('fichier', '')}").bold = True
            p.add_run(f" — {s.get('categorie', '')} (page {s.get('page', '?')})")
            if s.get("extrait"):
                ep = doc.add_paragraph()
                ep.add_run("   Extrait : ").italic = True
                ep.add_run(f'"{s["extrait"]}"').italic = True

    # ── Recommandations ────────────────────────────────────────
    recs = result.get("recommandations", [])
    if recs:
        doc.add_heading("4. RECOMMANDATIONS", level=1)
        prio_labels = {"IMMEDIATE": "⚡ IMMÉDIATE", "COURT_TERME": "📅 COURT TERME", "LONG_TERME": "📆 LONG TERME"}
        for i, rec in enumerate(recs, 1):
            rp = doc.add_paragraph()
            rp.add_run(f"R{i}. {rec.get('action', '')}").bold = True
            doc.add_paragraph(
                f"   Priorité : {prio_labels.get(rec.get('priorite', ''), rec.get('priorite', ''))} | "
                f"Responsable : {rec.get('responsable', 'Auditeur')}"
            )
            if rec.get("reference"):
                doc.add_paragraph(f"   Référence : {rec['reference']}")

    # ── Conclusion ─────────────────────────────────────────────
    doc.add_heading("5. CONCLUSION", level=1)
    doc.add_paragraph(result.get("conclusion", ""))

    # ── Documents analysés ─────────────────────────────────────
    if request.document_texts:
        doc.add_heading("6. DOCUMENTS ANALYSÉS", level=1)
        for doc_item in request.document_texts:
            dp = doc.add_paragraph(style='List Bullet')
            dp.add_run(doc_item.get("filename", "Document")).bold = True

    # ── Pied de page ───────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "─────────────────────────────────────────────\n"
        "Rapport généré par le système RAG Audit Juridique Marocain\n"
        "Base documentaire : Droit marocain (10 domaines juridiques)\n"
        "Ce rapport est fourni à titre indicatif. Consultez un juriste qualifié."
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Sauvegarde
    filename = f"rapport_rag_{audit_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = reports_dir / filename
    doc.save(str(path))
    print(f"✅ Rapport Word généré : {path}")
    return path


# ── Endpoints API ─────────────────────────────────────────────

@app.get("/health")
def health():
    vs = get_vectorstore()
    db_ready = vs is not None
    count = 0
    if db_ready:
        try:
            count = vs._collection.count()
        except Exception:
            pass
    return {
        "status": "ok",
        "vectorstore_ready": db_ready,
        "chunks_indexed": count,
        "models_available": list(AVAILABLE_MODELS.keys()),
        "docx_available": DOCX_AVAILABLE,
    }


@app.post("/analyse")
def start_analyse(request: AnalyseRequest, background_tasks: BackgroundTasks):
    audit_id = request.audit_id
    analyses_cache[audit_id] = AnalyseStatus(status="pending")
    background_tasks.add_task(run_rag_analysis, audit_id, request)
    return {"message": "Analyse lancée", "audit_id": audit_id, "status": "pending"}


@app.get("/analyse/{audit_id}/status")
def get_status(audit_id: str):
    if audit_id not in analyses_cache:
        raise HTTPException(status_code=404, detail="Analyse non trouvée")
    return analyses_cache[audit_id]


@app.get("/report/{audit_id}/download")
def download_report(audit_id: str):
    status = analyses_cache.get(audit_id)
    if not status or status.status != "done":
        raise HTTPException(status_code=404, detail="Rapport non disponible")
    report_path = status.result.get("report_path") if status.result else None
    if not report_path or not Path(report_path).exists():
        raise HTTPException(status_code=404, detail="Fichier rapport introuvable")
    return FileResponse(
        path=report_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=Path(report_path).name,
    )


if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*60)
    print("  ⚖️  RAG Audit Juridique Marocain — Service API")
    print("="*60)
    print(f"  Base vectorielle : {VECTOR_DB_PATH}")
    print(f"  Rapport Word     : {'✅ Disponible' if DOCX_AVAILABLE else '❌ Manquant (pip install python-docx)'}")
    print("  URL              : http://localhost:8000")
    print("  Docs             : http://localhost:8000/docs")
    print("="*60 + "\n")
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=False)
