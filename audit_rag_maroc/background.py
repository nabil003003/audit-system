# -*- coding: utf-8 -*-
"""Background jobs: Word report generation and volatile cache helpers."""

from __future__ import annotations

import logging
import threading
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None  # type: ignore

logger = logging.getLogger(__name__)

report_paths: dict[str, Path] = {}
_report_lock = threading.Lock()


def schedule_word_report(task_id: str, request: dict[str, Any], result: dict[str, Any]) -> None:
    """Run Word generation in a daemon thread so the HTTP worker returns immediately."""

    def _run() -> None:
        try:
            path = _generate_word_report(task_id, request, result)
            with _report_lock:
                report_paths[task_id] = path
            logger.info("Word report ready for task %s at %s", task_id, path)
        except Exception as exc:
            logger.error("Word report failed for task %s: %s", task_id, exc)

    threading.Thread(target=_run, name=f"word-{task_id[:8]}", daemon=True).start()


def take_report_path(task_id: str) -> Path | None:
    with _report_lock:
        return report_paths.pop(task_id, None)


def get_report_path(task_id: str) -> Path | None:
    with _report_lock:
        p = report_paths.get(task_id)
        return p if p and p.exists() else None


def _generate_word_report(task_id: str, request: dict[str, Any], result: dict[str, Any]) -> Path:
    if not DOCX_AVAILABLE or DocxDocument is None:
        raise RuntimeError("python-docx is not installed")
    reports_dir = Path(__file__).resolve().parent / "reports"
    reports_dir.mkdir(exist_ok=True)
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_par.add_run("LEGAL RAG AUDIT REPORT — MOROCCAN LAW")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Mission: {request.get('audit_title', '')}")
    sr.bold = True
    sr.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Analysis date: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        f"Model: {result.get('model_used', '')} | Task: {task_id[:12]}"
    )

    doc.add_page_break()
    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(result.get("summary", ""))

    analyse_globale = result.get("analyse_globale")
    if analyse_globale:
        doc.add_heading("2. Global Legal Analysis", level=1)
        doc.add_paragraph(analyse_globale)
        
    doc.add_heading("3. Non-conformities", level=1)
    violations = result.get("violations") or []
    if not violations:
        doc.add_paragraph("No major non-conformity detected in the supplied documents.")
    else:
        for i, v in enumerate(violations, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"Finding {i}: {v.get('titre', '')}")
            r.bold = True
            if v.get("texte_original"):
                doc.add_paragraph(v["texte_original"])
            if v.get("citation"):
                c = doc.add_paragraph()
                c.add_run(f"Legal quote: \"{v['citation']}\"").italic = True
            if v.get("source"):
                doc.add_paragraph(f"Source: {v['source']}")

    sources = result.get("sources") or []
    if sources:
        doc.add_heading("4. Legal sources used", level=1)
        for s in sources:
            doc.add_paragraph(f"{s.get('fichier', '')} — {s.get('categorie', '')}", style="List Bullet")

    recs = result.get("recommandations") or []
    if recs:
        doc.add_heading("5. Recommendations", level=1)
        for rec in recs:
            doc.add_paragraph(rec.get("action", ""), style="List Bullet")

    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(result.get("conclusion", ""))

    filename = f"rapport_rag_{task_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = reports_dir / filename
    doc.save(str(path))
    result["report_path"] = str(path)
    result["report_filename"] = path.name
    return path
